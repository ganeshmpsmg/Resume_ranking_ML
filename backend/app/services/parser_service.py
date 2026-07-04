"""
app/services/parser_service.py
---------------------------------
Resume Parser: extracts text from PDF/DOCX/TXT and detects structured
fields (name, contact info, sections, experience, education).
"""

from __future__ import annotations
import io
import re
import datetime
from dataclasses import dataclass, field

from app.core.taxonomy import EDUCATION_RANKS
from app.core.logging_config import get_logger
from app.utils.text_utils import (
    clean_text, EMAIL_REGEX, PHONE_REGEX, LINKEDIN_REGEX, GITHUB_REGEX,
    EXPERIENCE_YEARS_REGEX,
)

logger = get_logger(__name__)

try:
    import fitz  # PyMuPDF
    _HAS_FITZ = True
except ImportError:
    _HAS_FITZ = False

try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False

try:
    import docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


SECTION_HEADERS = {
    "summary": ["summary", "objective", "profile", "about me"],
    "education": ["education", "academic background", "qualifications"],
    "experience": ["experience", "work experience", "employment history",
                   "professional experience", "work history"],
    "skills": ["skills", "technical skills", "core competencies",
               "key skills", "skill set"],
    "projects": ["projects", "academic projects", "personal projects",
                 "key projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}


@dataclass
class ParsedResume:
    filename: str = ""
    raw_text: str = ""
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    linkedin: str | None = None
    github: str | None = None
    sections: dict = field(default_factory=dict)
    years_experience: float = 0.0
    education_level: str | None = None
    education_rank: int = 0


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    if _HAS_FITZ:
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                text = "\n".join(page.get_text("text") for page in doc)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF failed ({e}); falling back to pdfplumber")
    if _HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            return text
        except Exception as e:
            logger.error(f"pdfplumber also failed: {e}")
    if not _HAS_FITZ and not _HAS_PDFPLUMBER:
        raise RuntimeError("No PDF backend available (PyMuPDF / pdfplumber).")
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not _HAS_DOCX:
        raise RuntimeError("python-docx is not installed.")
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        raw = extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        raw = extract_text_from_docx(file_bytes)
    elif ext == "txt":
        raw = extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
    return clean_text(raw)


def detect_name(text: str, filename: str = "") -> str | None:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    name_pattern = re.compile(r"^([A-Z][a-zA-Z'.-]+(\s+[A-Z][a-zA-Z'.-]+){0,3})$")
    for line in lines[:8]:
        if EMAIL_REGEX.search(line) or LINKEDIN_REGEX.search(line):
            continue
        if any(ch.isdigit() for ch in line):
            continue
        if len(line.split()) > 5 or len(line) > 60:
            continue
        candidate = line.strip(" |•-:")
        if name_pattern.match(candidate):
            return candidate
    if filename:
        base = filename.rsplit(".", 1)[0]
        base = re.sub(r"(?i)resume|cv|_|-", " ", base)
        base = re.sub(r"\s+", " ", base).strip()
        if base and not any(ch.isdigit() for ch in base):
            return base.title()
    return None


def detect_contact_info(text: str) -> dict:
    email_match = EMAIL_REGEX.search(text)
    phone_match = PHONE_REGEX.search(text)
    linkedin_match = LINKEDIN_REGEX.search(text)
    github_match = GITHUB_REGEX.search(text)
    return {
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0).strip() if phone_match else None,
        "linkedin": linkedin_match.group(0) if linkedin_match else None,
        "github": github_match.group(0) if github_match else None,
    }


def segment_sections(text: str) -> dict:
    lines = text.split("\n")
    header_lookup = {}
    for section, keywords in SECTION_HEADERS.items():
        for kw in keywords:
            header_lookup[kw] = section

    sections = {"header": []}
    current_section = "header"
    for line in lines:
        stripped = line.strip().lower().strip(":-• ")
        matched_section = None
        if 0 < len(stripped) <= 40:
            for kw, section in header_lookup.items():
                if stripped == kw or stripped.startswith(kw):
                    matched_section = section
                    break
        if matched_section:
            current_section = matched_section
            sections.setdefault(current_section, [])
            continue
        sections.setdefault(current_section, []).append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


def estimate_years_experience(text: str) -> float:
    matches = EXPERIENCE_YEARS_REGEX.findall(text)
    if matches:
        try:
            years = [float(m) for m in matches if m]
            if years:
                return max(years)
        except ValueError:
            pass

    range_pattern = re.compile(
        r"(19|20)\d{2}\s*[-–to]+\s*((19|20)\d{2}|present|current)", re.IGNORECASE,
    )
    total_months = 0
    current_year = datetime.datetime.now().year
    for m in range_pattern.finditer(text):
        full = m.group(0)
        start_year = int(re.search(r"(19|20)\d{2}", full).group(0))
        if "present" in full.lower() or "current" in full.lower():
            end_year = current_year
        else:
            nums = [int(x) for x in re.findall(r"(?:19|20)\d{2}", full)]
            end_year = nums[-1] if len(nums) >= 2 else start_year
        span = max(0, end_year - start_year)
        total_months += span * 12
    return round(total_months / 12, 1)


def detect_education_level(text: str) -> tuple[str | None, int]:
    lowered = text.lower()
    best_level, best_rank = None, 0
    for keyword, rank in EDUCATION_RANKS.items():
        if keyword in lowered and rank > best_rank:
            best_level, best_rank = keyword, rank
    return best_level, best_rank


def parse_resume(filename: str, file_bytes: bytes) -> ParsedResume:
    raw_text = extract_text(filename, file_bytes)
    if not raw_text.strip():
        logger.warning(f"No text extracted from {filename}")

    contact = detect_contact_info(raw_text)
    name = detect_name(raw_text, filename)
    sections = segment_sections(raw_text)
    years_exp = estimate_years_experience(raw_text)
    edu_level, edu_rank = detect_education_level(raw_text)

    return ParsedResume(
        filename=filename,
        raw_text=raw_text,
        name=name,
        email=contact["email"],
        phone=contact["phone"],
        linkedin=contact["linkedin"],
        github=contact["github"],
        sections=sections,
        years_experience=years_exp,
        education_level=edu_level,
        education_rank=edu_rank,
    )
